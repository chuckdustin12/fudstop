
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

CONFIG = {
    "output_path": "caption_output.docx",
    "case_no": "322-744263-23",
    "petitioner": "MORGAN MICHELLE MYERS",
    "respondent": "CHARLES DUSTIN MYERS",
    "children_line": "M.E.M. AND C.R.M.,",
    "motion_title": "MOTION FOR NO-EVIDENCE SUMMARY JUDGMENT",
    "court_heading": "IN THE DISTRICT COURT",
    "judicial_district": "322ND JUDICIAL DISTRICT",
    "county_line": "TARRANT COUNTY, TEXAS",
    "debug_show_grid": True,   # set to False for only top/bottom rules once layout is confirmed
}

def _set_table_borders(table, grid=False):
    """If grid=True, draw full grid borders (for debugging). Else top/bottom only."""
    tblPr = table._tbl.tblPr
    borders = None
    for child in tblPr.iterchildren():
        if child.tag.endswith('tblBorders'):
            borders = child; break
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)

    def side(name, val, sz):
        el = None
        for c in borders.iterchildren():
            if c.tag.endswith(name):
                el = c; break
        if el is None:
            el = OxmlElement(f"w:{name}")
            borders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')

    if grid:
        for n in ('top','bottom','left','right','insideH','insideV'):
            side(n, 'single', 12)
    else:
        side('top','single',24)
        side('bottom','single',24)
        for n in ('left','right','insideH','insideV'):
            side(n, 'nil', 0)

def generate_caption_doc(
    path,
    case_no,
    petitioner,
    respondent,
    children_line,
    motion_title,
    court_heading,
    judicial_district,
    county_line,
    debug_show_grid=True,
):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Case number
    p = doc.add_paragraph(f"NO. {case_no}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for r in p.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    # Caption as ONE table (8 rows, 3 columns)
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    # Exact widths
    widths = (Inches(4.0), Inches(0.4), Inches(2.8))
    for j, w in enumerate(widths):
        for i in range(8):
            table.cell(i, j).width = w

    _set_table_borders(table, grid=debug_show_grid)

    left_lines = [
        "IN THE MATTER OF",
        "THE MARRIAGE OF",
        petitioner,
        "AND",
        respondent,
        "AND IN THE INTEREST OF",
        children_line,
        "CHILDREN",
    ]
    right_lines = [
        court_heading,
        "",
        judicial_district,
        county_line,
        "", "", "", "",
    ]

    for i in range(8):
        L = table.cell(i,0).paragraphs[0]
        M = table.cell(i,1).paragraphs[0]
        R = table.cell(i,2).paragraphs[0]

        L.text = left_lines[i]
        M.text = "§"
        R.text = right_lines[i]

        # spacing + fonts
        for par, align in ((L, WD_ALIGN_PARAGRAPH.LEFT), (M, WD_ALIGN_PARAGRAPH.CENTER), (R, WD_ALIGN_PARAGRAPH.LEFT)):
            par.alignment = align
            pf = par.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            for run in par.runs:
                run.font.name = "Times New Roman"; run.font.size = Pt(12)

        # vertical align
        table.cell(i,0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(i,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(i,2).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Title block with rules (using separate 1-row table for exact line)
    # top rule
    rule_top = doc.add_table(rows=1, cols=1)
    rule_top.allow_autofit = False
    rule_top.columns[0].width = Inches(7.2)
    _set_table_borders(rule_top, grid=True)  # 1-cell thick line

    # Title
    title = doc.add_paragraph(motion_title.upper())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for r in title.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12); r.font.bold = True; r.font.all_caps = True

    # bottom rule
    rule_bot = doc.add_table(rows=1, cols=1)
    rule_bot.allow_autofit = False
    rule_bot.columns[0].width = Inches(7.2)
    _set_table_borders(rule_bot, grid=True)

    # First body line
    p2 = doc.add_paragraph("TO THE HONORABLE JUDGE OF THE COURT:")
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for r in p2.runs:
        r.font.name = "Times New Roman"; r.font.size = Pt(12)

    doc.save(path)

def main():
    generate_caption_doc(
        path=CONFIG["output_path"],
        case_no=CONFIG["case_no"],
        petitioner=CONFIG["petitioner"],
        respondent=CONFIG["respondent"],
        children_line=CONFIG["children_line"],
        motion_title=CONFIG["motion_title"],
        court_heading=CONFIG["court_heading"],
        judicial_district=CONFIG["judicial_district"],
        county_line=CONFIG["county_line"],
        debug_show_grid=CONFIG["debug_show_grid"],
    )

main()
