
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn

CONFIG = {
    "output_path": "caption_output.docx",
    "case_no": "322-744263-23",
    "petitioner": "MORGAN MICHELLE MYERS",
    "respondent": "CHARLES DUSTIN MYERS",
    "children_line": "M.E.M. AND C.R.M.,",
    "motion_title": "MOTION FOR NO-EVIDENCE SUMMARY JUDGMENT",
    "court_heading": "IN THE DISTRICT COURT",
    "judicial_district": "322ND JUDICIAL DISTRICT",
    "county_line": "TARRANT COUNTY, TEXAS",
    "show_bottom_line": True,
}

def _set_cell_margins(cell, top=80, left=80, bottom=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # find or create tcMar
    tcMar = None
    for child in tcPr.iterchildren():
        if child.tag.endswith('tcMar'):
            tcMar = child
            break
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    def set_side(tag, val):
        el = None
        for child in tcMar.iterchildren():
            if child.tag.endswith(tag):
                el = child
                break
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
    set_side('top', top)
    set_side('left', left)
    set_side('bottom', bottom)
    set_side('right', right)

def _set_paragraph_border(p, top=False, bottom=False):
    pPr = p._p.get_or_add_pPr()
    pbdr = None
    for child in pPr.iterchildren():
        if child.tag.endswith('pBdr'):
            pbdr = child
            break
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        pPr.append(pbdr)
    def set_side(tag):
        el = None
        for child in pbdr.iterchildren():
            if child.tag.endswith(tag):
                el = child
                break
        if el is None:
            el = OxmlElement(f'w:{tag}')
            pbdr.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12')
        el.set(qn('w:space'), '1')
        el.set(qn('w:color'), '000000')
    if top: set_side('top')
    if bottom: set_side('bottom')

def _table_top_bottom_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = None
    for child in tblPr.iterchildren():
        if child.tag.endswith('tblBorders'):
            borders = child
            break
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    def set_tbl_side(name, sz, val):
        el = None
        for child in borders.iterchildren():
            if child.tag.endswith(name):
                el = child
                break
        if el is None:
            el = OxmlElement(f'w:{name}')
            borders.append(el)
        el.set(qn('w:val'), val)
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
    set_tbl_side('top', 16, 'single')
    set_tbl_side('bottom', 16, 'single')
    for name in ('left','right','insideH','insideV'):
        set_tbl_side(name, 0, 'nil')

def generate_caption_doc(
    path,
    case_no,
    petitioner,
    respondent,
    children_line,
    motion_title,
    court_heading,
    judicial_district,
    county_line,
    show_bottom_line=True,
):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    p = doc.add_paragraph(f"NO. {case_no}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _table_top_bottom_borders(table)

    left_lines = [
        "IN THE MATTER OF",
        "THE MARRIAGE OF",
        petitioner,
        "AND",
        respondent,
        "AND IN THE INTEREST OF",
        children_line,
        "CHILDREN",
    ]
    right_lines = [
        court_heading,
        "",
        judicial_district,
        county_line,
        "", "", "", "",
    ]

    for i in range(8):
        row = table.rows[i]
        row.cells[0].text = left_lines[i]
        row.cells[1].text = "§"
        row.cells[2].text = right_lines[i]
        for j, cell in enumerate(row.cells):
            _set_cell_margins(cell, top=20, left=40, bottom=20, right=40)
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Border rule above title
    spacer = doc.add_paragraph("")
    _set_paragraph_border(spacer, bottom=True)

    title = doc.add_paragraph(motion_title.upper())
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.all_caps = True

    if show_bottom_line:
        spacer2 = doc.add_paragraph("")
        _set_paragraph_border(spacer2, top=True)

    p2 = doc.add_paragraph("TO THE HONORABLE JUDGE OF THE COURT:")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p2.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.save(path)

def main():
    generate_caption_doc(
        path=CONFIG["output_path"],
        case_no=CONFIG["case_no"],
        petitioner=CONFIG["petitioner"],
        respondent=CONFIG["respondent"],
        children_line=CONFIG["children_line"],
        motion_title=CONFIG["motion_title"],
        court_heading=CONFIG["court_heading"],
        judicial_district=CONFIG["judicial_district"],
        county_line=CONFIG["county_line"],
        show_bottom_line=CONFIG["show_bottom_line"],
    )

